import os, io, re, secrets
from datetime import datetime, timezone, timedelta
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Request
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse
from sqlalchemy.orm import Session
from typing import List

from database import engine, Base, get_db
from models import (
    Organization, User, Blueprint, BlueprintState, Template, TemplateBlueprint,
    Project, Entry, EntryImage, Comment, InviteLink, Tag,
)
from schemas import (
    OrganizationCreate, OrganizationOut,
    UserOut,
    BlueprintCreate, BlueprintOut, BlueprintStateCreate, BlueprintStateOut,
    TemplateCreate, TemplateOut,
    ProjectCreate, ProjectOut, ProjectUpdate,
    EntryCreate, EntryUpdate, EntryOut,
    EntryImageOut,
    CommentCreate, CommentOut,
    InviteLinkCreate, InviteLinkOut,
    TagOut, TagCreate,
)

Base.metadata.create_all(bind=engine)

BASE_DIR = os.path.dirname(__file__)
UPLOAD_DIR = os.path.join(BASE_DIR, "..", "uploads")
FRONTEND_DIR = os.path.join(BASE_DIR, "..", "frontend")
REEL_SPEC_PATH = os.path.join(BASE_DIR, "..", "reel-spec.html")
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(title="Choreography Manager")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")
app.mount("/app", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")


@app.get("/reel-spec")
def serve_reel_spec():
    from fastapi.responses import FileResponse
    return FileResponse(REEL_SPEC_PATH)


@app.get("/api/health")
def health():
    return {"ok": True}


# ── DEPENDENCIES ──

from config import GOOGLE_CLIENT_ID

def get_dev_user(db: Session):
    """Get or create dev user for local development (no Google OAuth)."""
    user = db.query(User).filter(User.google_sub == "dev-local").first()
    if not user:
        org = db.query(Organization).filter(Organization.id == 1).first()
        if not org:
            org = Organization(id=1, name="Internal", plan="internal")
            db.add(org)
            db.commit()
        user = User(
            organization_id=1, email="dev@localhost", name="Dev User",
            google_sub="dev-local", role="admin", approved=True,
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    return user


def get_current_user(request: Request, db: Session = Depends(get_db)):
    # Dev mode: no Google OAuth configured, auto-login as dev user
    if not GOOGLE_CLIENT_ID:
        return get_dev_user(db)
    user_id = request.cookies.get("user_id")
    if not user_id:
        raise HTTPException(401, "Not logged in")
    user = db.query(User).filter(User.id == int(user_id)).first()
    if not user:
        raise HTTPException(401, "User not found")
    if not user.approved:
        raise HTTPException(403, "Account pending approval")
    return user


def require_admin(user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Admin required")
    return user




# ── AUTH ──

@app.get("/auth/login")
def auth_login():
    from config import GOOGLE_CLIENT_ID
    redirect_uri = "http://localhost:8000/auth/callback"
    return RedirectResponse(
        f"https://accounts.google.com/o/oauth2/v2/auth?"
        f"client_id={GOOGLE_CLIENT_ID}&redirect_uri={redirect_uri}"
        f"&response_type=code&scope=openid email profile&access_type=offline"
    )


@app.get("/auth/callback")
def auth_callback(code: str = "", db: Session = Depends(get_db)):
    if not code:
        raise HTTPException(400, "No code provided")
    # Exchange code for tokens (simplified - in production use httpx)
    # For now, this is a stub that needs Google credentials configured
    return {"status": "callback received", "code": code[:10] + "..."}


@app.post("/auth/login-token")
def auth_login_token(token: str, db: Session = Depends(get_db)):
    # Verify Google ID token and get user info
    # Stub: in production, verify against Google's public keys
    from config import GOOGLE_CLIENT_ID
    from authlib.integrations.starlette_client import OAuth
    # This is a placeholder - real implementation verifies the JWT
    raise HTTPException(501, "Configure Google OAuth first")


@app.get("/api/me")
def get_me(user: User = Depends(get_current_user)):
    return user


@app.post("/auth/logout")
def auth_logout():
    response = RedirectResponse("/")
    response.delete_cookie("user_id")
    return response


# ── USERS (admin) ──

@app.get("/api/users/pending", response_model=List[UserOut])
def list_pending_users(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    return db.query(User).filter(User.approved == False).all()


@app.post("/api/users/{user_id}/approve")
def approve_user(user_id: int, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(404)
    user.approved = True
    db.commit()
    return {"ok": True}


@app.post("/api/users/{user_id}/reject")
def reject_user(user_id: int, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(404)
    db.delete(user)
    db.commit()
    return {"ok": True}


# ── INVITE LINKS ──

@app.post("/api/organizations/{org_id}/invites", response_model=InviteLinkOut)
def create_invite(org_id: int, data: InviteLinkCreate, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    if admin.organization_id != org_id:
        raise HTTPException(403)
    token = secrets.token_urlsafe(32)
    invite = InviteLink(
        organization_id=org_id,
        token=token,
        email_optional=data.email_optional,
        expires_at=datetime.now(timezone.utc) + timedelta(days=7),
    )
    db.add(invite)
    db.commit()
    db.refresh(invite)
    return invite


@app.get("/api/invites/{token}")
def validate_invite(token: str, db: Session = Depends(get_db)):
    invite = db.query(InviteLink).filter(InviteLink.token == token).first()
    if not invite:
        raise HTTPException(404, "Invite not found")
    if invite.used_at:
        raise HTTPException(400, "Invite already used")
    if invite.expires_at < datetime.now(timezone.utc):
        raise HTTPException(400, "Invite expired")
    return {"valid": True, "organization_id": invite.organization_id}


# ── ORGANIZATIONS ──

@app.get("/api/organizations", response_model=List[OrganizationOut])
def list_organizations(db: Session = Depends(get_db)):
    return db.query(Organization).all()


# ── BLUEPRINTS ──

@app.post("/api/blueprints", response_model=BlueprintOut)
def create_blueprint(bp: BlueprintCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    existing = db.query(Blueprint).filter(Blueprint.name == bp.name, Blueprint.organization_id == user.organization_id).first()
    if existing:
        raise HTTPException(400, "Blueprint name already exists")
    obj = Blueprint(name=bp.name, description=bp.description, organization_id=user.organization_id)
    for s in bp.states:
        obj.states.append(BlueprintState(**s.model_dump()))
    db.add(obj)
    db.commit()
    db.refresh(obj)
    return obj


@app.get("/api/blueprints", response_model=List[BlueprintOut])
def list_blueprints(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Blueprint).filter(Blueprint.organization_id == user.organization_id).order_by(Blueprint.name).all()


@app.get("/api/blueprints/{bp_id}", response_model=BlueprintOut)
def get_blueprint(bp_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Blueprint).filter(Blueprint.id == bp_id, Blueprint.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    return obj


@app.put("/api/blueprints/{bp_id}", response_model=BlueprintOut)
def update_blueprint(bp_id: int, bp: BlueprintCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Blueprint).filter(Blueprint.id == bp_id, Blueprint.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    obj.name = bp.name
    obj.description = bp.description
    db.query(BlueprintState).filter(BlueprintState.blueprint_id == bp_id).delete()
    for s in bp.states:
        db.add(BlueprintState(blueprint_id=bp_id, **s.model_dump()))
    db.commit()
    db.refresh(obj)
    return obj


@app.delete("/api/blueprints/{bp_id}")
def delete_blueprint(bp_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Blueprint).filter(Blueprint.id == bp_id, Blueprint.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    db.delete(obj)
    db.commit()
    return {"ok": True}


# ── TEMPLATES ──

@app.post("/api/templates", response_model=TemplateOut)
def create_template(t: TemplateCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = Template(name=t.name, description=t.description, organization_id=user.organization_id)
    for i, bp_id in enumerate(t.blueprint_ids):
        obj.blueprints.append(TemplateBlueprint(blueprint_id=bp_id, sort_order=i))
    db.add(obj)
    db.commit()
    db.refresh(obj)
    return obj


@app.get("/api/templates", response_model=List[TemplateOut])
def list_templates(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Template).filter(Template.organization_id == user.organization_id).order_by(Template.name).all()


@app.get("/api/templates/{t_id}", response_model=TemplateOut)
def get_template(t_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Template).filter(Template.id == t_id, Template.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    return obj


@app.put("/api/templates/{t_id}", response_model=TemplateOut)
def update_template(t_id: int, t: TemplateCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Template).filter(Template.id == t_id, Template.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    obj.name = t.name
    obj.description = t.description
    db.query(TemplateBlueprint).filter(TemplateBlueprint.template_id == t_id).delete()
    for i, bp_id in enumerate(t.blueprint_ids):
        db.add(TemplateBlueprint(template_id=t_id, blueprint_id=bp_id, sort_order=i))
    db.commit()
    db.refresh(obj)
    return obj


@app.delete("/api/templates/{t_id}")
def delete_template(t_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Template).filter(Template.id == t_id, Template.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    db.delete(obj)
    db.commit()
    return {"ok": True}


# ── PROJECTS ──

@app.post("/api/projects", response_model=ProjectOut)
def create_project(p: ProjectCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = Project(
        name=p.name, template_id=p.template_id, organization_id=user.organization_id,
        game_type=p.game_type, customer=p.customer, deadline=p.deadline, asset_link=p.asset_link,
    )
    db.add(obj)
    db.flush()

    if p.template_id:
        template = db.query(Template).filter(Template.id == p.template_id).first()
        if template:
            for tb in template.blueprints:
                bp = tb.blueprint
                if not bp:
                    continue
                for state in bp.states:
                    db.add(Entry(
                        project_id=obj.id,
                        element_name=bp.name,
                        animation_name=state.name,
                        looping=state.default_looping,
                        duration=state.default_duration,
                        description=state.default_description,
                        projected_hours=0.0,
                        actual_hours=0.0,
                    ))
    db.commit()
    db.refresh(obj)
    return obj


@app.get("/api/projects", response_model=List[ProjectOut])
def list_projects(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Project).filter(Project.organization_id == user.organization_id).order_by(Project.created_at.desc()).all()


@app.get("/api/projects/{p_id}", response_model=ProjectOut)
def get_project(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    return obj


@app.put("/api/projects/{p_id}")
def update_project(p_id: int, data: ProjectUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    for k, v in data.model_dump(exclude_unset=True).items():
        setattr(obj, k, v)
    db.commit()
    return {"ok": True}


@app.delete("/api/projects/{p_id}")
def delete_project(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not obj:
        raise HTTPException(404)
    db.delete(obj)
    db.commit()
    return {"ok": True}


# ── ENTRIES ──

@app.get("/api/projects/{p_id}/entries", response_model=List[EntryOut])
def list_entries(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    entries = db.query(Entry).filter(Entry.project_id == p_id).order_by(Entry.element_name, Entry.id).all()
    return [EntryOut.model_validate(e) for e in entries]


@app.get("/api/entries/{e_id}", response_model=EntryOut)
def get_entry(e_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    return EntryOut.model_validate(obj)


@app.post("/api/entries", response_model=EntryOut)
def create_entry(e: EntryCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == e.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    obj = Entry(**e.model_dump())
    db.add(obj)
    db.commit()
    db.refresh(obj)
    return EntryOut.model_validate(obj)


@app.put("/api/entries/{e_id}", response_model=EntryOut)
def update_entry(e_id: int, e: EntryUpdate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == obj.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    update_data = e.model_dump(exclude_unset=True)
    for k, v in update_data.items():
        setattr(obj, k, v)
    if "actual_hours" in update_data or "projected_hours" in update_data:
        from config import ALERT_THRESHOLD
        if obj.projected_hours > 0 and obj.actual_hours > obj.projected_hours * ALERT_THRESHOLD:
            obj.alert_flag = True
            obj.alert_flag_reason = f"auto: {obj.actual_hours}h > {obj.projected_hours}h * {ALERT_THRESHOLD}"
    db.commit()
    db.refresh(obj)
    return EntryOut.model_validate(obj)


@app.delete("/api/entries/{e_id}")
def delete_entry(e_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == obj.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    db.delete(obj)
    db.commit()
    return {"ok": True}


@app.post("/api/entries/{e_id}/flag")
def toggle_entry_flag(e_id: int, data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == obj.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    obj.alert_flag = data.get("alert_flag", not obj.alert_flag)
    obj.alert_flag_reason = data.get("alert_flag_reason", f"manual: flagged by {user.name}")
    db.commit()
    return {"ok": True}


# ── ENTRY IMAGES ──

@app.get("/api/entries/{e_id}/images", response_model=List[EntryImageOut])
def list_entry_images(e_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == obj.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    return db.query(EntryImage).filter(EntryImage.entry_id == e_id).order_by(EntryImage.sort_order).all()


@app.post("/api/entries/{e_id}/images", response_model=EntryImageOut)
def upload_entry_image(e_id: int, file: UploadFile = File(...), user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == obj.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)

    ext = os.path.splitext(file.filename)[1] if file.filename else ".png"
    filename = f"entry_{e_id}_{int(datetime.now(timezone.utc).timestamp())}{ext}"
    path = os.path.join(UPLOAD_DIR, filename)

    with open(path, "wb") as f:
        f.write(file.file.read())

    # WebP conversion
    try:
        from PIL import Image
        img = Image.open(path)
        webp_path = path.rsplit(".", 1)[0] + ".webp"
        img.save(webp_path, "WEBP", quality=80)
        os.remove(path)
        filename = os.path.basename(webp_path)
    except Exception:
        pass  # Keep original if conversion fails

    max_order = db.query(EntryImage).filter(EntryImage.entry_id == e_id).count()
    image = EntryImage(entry_id=e_id, image_path=filename, sort_order=max_order)
    db.add(image)
    db.commit()
    db.refresh(image)
    return image


@app.delete("/api/entries/{e_id}/images/{image_id}")
def delete_entry_image(e_id: int, image_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    image = db.query(EntryImage).filter(EntryImage.id == image_id, EntryImage.entry_id == e_id).first()
    if not image:
        raise HTTPException(404)
    path = os.path.join(UPLOAD_DIR, image.image_path)
    if os.path.exists(path):
        os.remove(path)
    db.delete(image)
    db.commit()
    return {"ok": True}


@app.put("/api/entries/{e_id}/images/reorder")
def reorder_entry_images(e_id: int, data: dict, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    order = data.get("order", [])
    for idx, img_id in enumerate(order):
        img = db.query(EntryImage).filter(EntryImage.id == img_id, EntryImage.entry_id == e_id).first()
        if img:
            img.sort_order = idx
    db.commit()
    return {"ok": True}


# ── LEGACY IMAGE UPLOAD (kept for backwards compat) ──

@app.post("/api/entries/{e_id}/image")
def upload_entry_image_legacy(e_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj:
        raise HTTPException(404)
    ext = os.path.splitext(file.filename)[1] if file.filename else ".png"
    filename = f"entry_{e_id}_{int(datetime.now(timezone.utc).timestamp())}{ext}"
    path = os.path.join(UPLOAD_DIR, filename)
    with open(path, "wb") as f:
        f.write(file.file.read())
    obj.image_path = filename
    db.commit()
    return {"image_path": filename}


@app.delete("/api/entries/{e_id}/image")
def delete_entry_image_legacy(e_id: int, db: Session = Depends(get_db)):
    obj = db.query(Entry).filter(Entry.id == e_id).first()
    if not obj or not obj.image_path:
        raise HTTPException(404)
    path = os.path.join(UPLOAD_DIR, obj.image_path)
    if os.path.exists(path):
        os.remove(path)
    obj.image_path = ""
    db.commit()
    return {"ok": True}


# ── COMMENTS ──

@app.post("/api/projects/{p_id}/comments", response_model=CommentOut)
def create_comment(p_id: int, data: CommentCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    comment = Comment(
        project_id=p_id, entry_id=data.entry_id, author_id=user.id,
        body=data.body, linked_comment_id=data.linked_comment_id,
    )
    db.add(comment)
    db.commit()
    db.refresh(comment)
    return comment


@app.get("/api/projects/{p_id}/comments", response_model=List[CommentOut])
def list_comments(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    return db.query(Comment).filter(Comment.project_id == p_id).order_by(Comment.created_at).all()


@app.delete("/api/comments/{c_id}")
def delete_comment(c_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    comment = db.query(Comment).filter(Comment.id == c_id).first()
    if not comment:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == comment.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    if comment.author_id != user.id and user.role != "admin":
        raise HTTPException(403)
    db.delete(comment)
    db.commit()
    return {"ok": True}


# ── TAGS ──

@app.get("/api/projects/{p_id}/tags", response_model=List[TagOut])
def list_project_tags(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    return db.query(Tag).filter(Tag.project_id == p_id).order_by(Tag.name).all()


@app.post("/api/projects/{p_id}/tags", response_model=TagOut)
def create_project_tag(p_id: int, data: TagCreate, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    existing = db.query(Tag).filter(Tag.project_id == p_id, Tag.name == data.name).first()
    if existing:
        raise HTTPException(400, detail="Tag already exists")
    tag = Tag(project_id=p_id, name=data.name)
    db.add(tag)
    db.commit()
    db.refresh(tag)
    return tag


@app.delete("/api/tags/{tag_id}")
def delete_tag(tag_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    tag = db.query(Tag).filter(Tag.id == tag_id).first()
    if not tag:
        raise HTTPException(404)
    project = db.query(Project).filter(Project.id == tag.project_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(403)
    db.delete(tag)
    db.commit()
    return {"ok": True}


# ── ROLLUP ──

@app.get("/api/projects/{p_id}/rollup")
def project_rollup(p_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    entries = db.query(Entry).filter(Entry.project_id == p_id).all()
    total_projected = sum(e.projected_hours or 0 for e in entries)
    total_actual = sum(e.actual_hours or 0 for e in entries)
    by_element = {}
    by_artist = {}
    for e in entries:
        ph = e.projected_hours or 0
        ah = e.actual_hours or 0
        by_element.setdefault(e.element_name, {"element": e.element_name, "projected": 0, "actual": 0, "count": 0})
        by_element[e.element_name]["projected"] += ph
        by_element[e.element_name]["actual"] += ah
        by_element[e.element_name]["count"] += 1
        if e.artist:
            by_artist.setdefault(e.artist, {"artist": e.artist, "projected": 0, "actual": 0, "count": 0})
            by_artist[e.artist]["projected"] += ph
            by_artist[e.artist]["actual"] += ah
            by_artist[e.artist]["count"] += 1
    flagged = [e.id for e in entries if e.alert_flag]
    return {
        "total_projected": total_projected,
        "total_actual": total_actual,
        "total_entries": len(entries),
        "flagged_count": len(flagged),
        "by_element": sorted(by_element.values(), key=lambda x: -x["actual"]),
        "by_artist": sorted(by_artist.values(), key=lambda x: -x["actual"]),
    }


# ── REEL SPEC IMAGE UPLOAD ──

@app.post("/api/reel/upload")
def upload_reel_image(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename)[1] if file.filename else ".png"
    filename = f"reel_{int(datetime.now(timezone.utc).timestamp())}_{os.urandom(4).hex()}{ext}"
    path = os.path.join(UPLOAD_DIR, filename)
    with open(path, "wb") as f:
        f.write(file.file.read())
    return {"image_path": filename}


@app.delete("/api/reel/upload")
def delete_reel_image(data: dict):
    filename = data.get("image_path", "")
    if not filename:
        raise HTTPException(400, "No image_path provided")
    path = os.path.join(UPLOAD_DIR, filename)
    if os.path.exists(path):
        os.remove(path)
    return {"ok": True}


# ── IMPORT DOCX ──

@app.post("/api/import-docx")
def import_docx(file: UploadFile = File(...)):
    from docx import Document
    data = file.file.read()
    doc = Document(io.BytesIO(data))
    entries = []
    current_element = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text or len(text) < 1:
            continue
        lower = text.lower()
        if lower.startswith("summary") or lower.startswith("design") or lower.startswith("animation total"):
            continue
        if len(text) < 30 and not re.search(r'\bhr\b', lower) and not re.search(r'\d', text):
            current_element = text
        hr_match = re.findall(r'(\d+[\d.]*)\s*hr', lower)
        if hr_match and current_element:
            total_h = sum(float(h) for h in hr_match)
            anim_name = re.sub(r'\s*:?\s*\d+[\d.]*\s*hr.*$', '', text).strip()
            if anim_name and len(anim_name) > 1 and not anim_name.startswith(" ") and current_element:
                entries.append({
                    "element_name": current_element, "animation_name": anim_name,
                    "looping": "loop" in lower or "idle" in lower,
                    "duration": "", "description": "", "artist": "",
                    "projected_hours": total_h, "actual_hours": 0.0,
                })

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 4:
            continue
        label0 = rows[0].cells[0].text.strip().lower()
        if "symbol" not in label0 and "name" not in label0:
            continue
        looping, duration, description, artist = "", "", "", ""
        for row in rows:
            cells = [c.text.strip() for c in row.cells]
            label = cells[0].lower() if cells else ""
            val = cells[1] if len(cells) > 1 else ""
            meta = cells[-1] if len(cells) > 2 else ""
            if "loop" in label:
                looping = val
            elif "dur" in label:
                duration = val
            elif "desc" in label:
                description = val
            elif "symbol" in label or "name" in label:
                artist = meta or val
        hr_match = re.findall(r'(\d+[\d.]*)\s*hr', (artist + " " + description).lower())
        projected = sum(float(h) for h in hr_match) if hr_match else 0.0
        if description or duration:
            entries.append({
                "element_name": current_element or "", "animation_name": f"{current_element or ''}",
                "looping": looping.lower() in ("yes", "y", "true"),
                "duration": duration, "description": description,
                "artist": re.sub(r'\s*\d+[\d.]*\s*hr.*$', '', artist).strip() if artist else "",
                "projected_hours": projected, "actual_hours": 0.0,
            })
    return entries


# ── EXPORT ──

@app.get("/api/projects/{p_id}/export")
def export_project(p_id: int, format: str = "xlsx", user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == p_id, Project.organization_id == user.organization_id).first()
    if not project:
        raise HTTPException(404)
    entries = db.query(Entry).filter(Entry.project_id == p_id).order_by(Entry.element_name, Entry.id).all()
    total_hours = sum(e.actual_hours or 0 for e in entries)

    if format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()
        doc.add_heading(project.name, level=1)
        doc.add_paragraph("Choreography")
        doc.add_heading(f"Animation list (total {total_hours:.0f} hr)", level=2)

        # Group entries by element
        elements = {}
        for e in entries:
            elements.setdefault(e.element_name, []).append(e)

        # Summary table — grouped by element, merged element cells
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cell_shading(cell, color_hex):
            """Set background color on a table cell."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color_hex)
            shading.set(qn('w:val'), 'clear')
            tcPr.append(shading)

        summary_table = doc.add_table(rows=1, cols=3)
        summary_table.style = 'Table Grid'
        hdr = summary_table.rows[0].cells
        hdr[0].text = "Element"
        hdr[1].text = "Animation"
        hdr[2].text = "Hours"
        for cell in hdr:
            set_cell_shading(cell, "B4C6E7")  # light blue
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        # Track merge ranges and total rows
        merge_ranges = []
        row_idx = 1  # start after header
        for element_name, elem_entries in elements.items():
            start_row = row_idx
            elem_total = 0
            for i, e in enumerate(elem_entries):
                row = summary_table.add_row().cells
                row[0].text = element_name if i == 0 else ""
                row[1].text = e.animation_name or "Untitled"
                row[2].text = f"{e.actual_hours or 0:.1f}"
                elem_total += e.actual_hours or 0
                row_idx += 1
            # Add total row (included in merge)
            total_row = summary_table.add_row().cells
            total_row[0].text = ""
            total_row[1].text = f"{element_name} Total"
            total_row[2].text = f"{elem_total:.1f}"
            set_cell_shading(total_row[1], "D9E2F3")
            set_cell_shading(total_row[2], "D9E2F3")
            for p in total_row[1].paragraphs:
                for r in p.runs:
                    r.bold = True
            for p in total_row[2].paragraphs:
                for r in p.runs:
                    r.bold = True
            row_idx += 1
            # Merge includes all rows: entries + total row
            merge_ranges.append((start_row, row_idx - 1))

        # Merge element cells vertically (always, including single-entry elements)
        for start, end in merge_ranges:
            summary_table.cell(start, 0).merge(summary_table.cell(end, 0))

        doc.add_paragraph("")

        # Per-element sections
        for element_name, elem_entries in elements.items():
            doc.add_heading(element_name, level=2)
            for entry in elem_entries:
                doc.add_heading(entry.animation_name or "Untitled", level=3)
                tbl = doc.add_table(rows=4, cols=2)
                tbl.style = 'Table Grid'

                # Symbol row — image or text
                tbl.rows[0].cells[0].text = "Symbol"
                for r in tbl.rows[0].cells[0].paragraphs[0].runs:
                    r.bold = True
                img_cell = tbl.rows[0].cells[1]
                images = db.query(EntryImage).filter(EntryImage.entry_id == entry.id).order_by(EntryImage.sort_order).first()
                if images:
                    img_path = os.path.join(UPLOAD_DIR, images.image_path)
                    if os.path.exists(img_path):
                        try:
                            from PIL import Image as PilImage
                            pil_img = PilImage.open(img_path)
                            if pil_img.mode in ("RGBA", "P"):
                                pil_img = pil_img.convert("RGB")
                            png_buf = io.BytesIO()
                            pil_img.save(png_buf, "PNG")
                            png_buf.seek(0)
                            p = img_cell.paragraphs[0]
                            run = p.add_run()
                            run.add_picture(png_buf, width=Inches(2))
                        except Exception:
                            img_cell.text = entry.element_name
                    else:
                        img_cell.text = entry.element_name
                else:
                    img_cell.text = entry.element_name

                # Other rows
                detail_rows = [
                    ("Looping", "yes" if entry.looping else "no"),
                    ("Duration", entry.duration or ""),
                    ("Description", entry.description or ""),
                ]
                for i, (lbl, val) in enumerate(detail_rows, 1):
                    tbl.rows[i].cells[0].text = lbl
                    tbl.rows[i].cells[1].text = val
                    for r in tbl.rows[i].cells[0].paragraphs[0].runs:
                        r.bold = True
                doc.add_paragraph("")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from fastapi.responses import StreamingResponse
        safe_name = re.sub(r'[^\w\s-]', '', project.name).strip().replace(' ', '_')
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_choreography.docx"'},
        )

    if format == "csv":
        import csv, io
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Element", "Animation", "Looping", "Duration", "Description", "Artist", "Phase", "Status", "Priority", "Projected Hours", "Actual Hours", "Flag"])
        total_projected = 0
        total_actual = 0
        for e in entries:
            writer.writerow([
                e.element_name, e.animation_name, "Yes" if e.looping else "No",
                e.duration, e.description, e.artist, e.phase, e.status,
                e.priority, e.projected_hours or 0, e.actual_hours or 0,
                "Flagged" if e.alert_flag else "",
            ])
            total_projected += e.projected_hours or 0
            total_actual += e.actual_hours or 0
        writer.writerow(["TOTAL", "", "", "", "", "", "", "", "", f"{total_projected:.1f}", f"{total_actual:.1f}", ""])
        buf.seek(0)
        from fastapi.responses import StreamingResponse
        safe_name = re.sub(r'[^\w\s-]', '', project.name).strip().replace(' ', '_')
        return StreamingResponse(
            iter([buf.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_export.csv"'},
        )

    # Default: Excel
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    import io

    wb = Workbook()

    # Summary sheet
    ws = wb.active
    ws.title = "Summary"
    header_font = Font(bold=True, size=12)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, color="FFFFFF", size=11)

    ws["A1"] = project.name
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Choreography"
    ws["A3"] = f"Total Hours: {total_hours:.1f}"
    ws["A4"] = f"Projected Hours: {sum(e.projected_hours or 0 for e in entries):.1f}"
    if project.game_type:
        ws["B2"] = f"Game Type: {project.game_type}"
    if project.customer:
        ws["B3"] = f"Customer: {project.customer}"
    if project.deadline:
        ws["B4"] = f"Deadline: {project.deadline}"

    # Entries sheet
    ws2 = wb.create_sheet("Entries")
    headers = ["Element", "Animation", "Looping", "Duration", "Description", "Artist", "Phase", "Status", "Priority", "Projected Hours", "Actual Hours", "Flag"]
    for col, h in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_idx, e in enumerate(entries, 2):
        ws2.cell(row=row_idx, column=1, value=e.element_name)
        ws2.cell(row=row_idx, column=2, value=e.animation_name)
        ws2.cell(row=row_idx, column=3, value="Yes" if e.looping else "No")
        ws2.cell(row=row_idx, column=4, value=e.duration)
        ws2.cell(row=row_idx, column=5, value=e.description)
        ws2.cell(row=row_idx, column=6, value=e.artist)
        ws2.cell(row=row_idx, column=7, value=e.phase)
        ws2.cell(row=row_idx, column=8, value=e.status)
        ws2.cell(row=row_idx, column=9, value=e.priority)
        ws2.cell(row=row_idx, column=10, value=e.projected_hours or 0)
        ws2.cell(row=row_idx, column=11, value=e.actual_hours or 0)
        ws2.cell(row=row_idx, column=12, value="Flagged" if e.alert_flag else "")

    # Total row
    total_row = len(entries) + 2
    total_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
    ws2.cell(row=total_row, column=1, value="TOTAL").font = Font(bold=True)
    ws2.cell(row=total_row, column=1).fill = total_fill
    ws2.cell(row=total_row, column=10, value=sum(e.projected_hours or 0 for e in entries)).font = Font(bold=True)
    ws2.cell(row=total_row, column=10).fill = total_fill
    ws2.cell(row=total_row, column=11, value=sum(e.actual_hours or 0 for e in entries)).font = Font(bold=True)
    ws2.cell(row=total_row, column=11).fill = total_fill
    for col_idx in range(2, 13):
        ws2.cell(row=total_row, column=col_idx).fill = total_fill

    for col in ws2.columns:
        max_len = max(len(str(c.value or "")) for c in col)
        ws2.column_dimensions[col[0].column_letter].width = min(max_len + 2, 40)

    # Rollup sheet
    ws3 = wb.create_sheet("Rollup")
    ws3["A1"] = "Element"
    ws3["B1"] = "Projected"
    ws3["C1"] = "Actual"
    ws3["D1"] = "Entries"
    for cell in ws3[1]:
        cell.font = header_font_white
        cell.fill = header_fill

    by_element = {}
    for e in entries:
        by_element.setdefault(e.element_name, {"projected": 0, "actual": 0, "count": 0})
        by_element[e.element_name]["projected"] += e.projected_hours or 0
        by_element[e.element_name]["actual"] += e.actual_hours or 0
        by_element[e.element_name]["count"] += 1

    for row_idx, (name, data) in enumerate(sorted(by_element.items(), key=lambda x: -x[1]["actual"]), 2):
        ws3.cell(row=row_idx, column=1, value=name)
        ws3.cell(row=row_idx, column=2, value=data["projected"])
        ws3.cell(row=row_idx, column=3, value=data["actual"])
        ws3.cell(row=row_idx, column=4, value=data["count"])

    for col in ws3.columns:
        ws3.column_dimensions[col[0].column_letter].width = 15

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    safe_name = re.sub(r'[^\w\s-]', '', project.name).strip().replace(' ', '_')
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_export.xlsx"'},
    )
